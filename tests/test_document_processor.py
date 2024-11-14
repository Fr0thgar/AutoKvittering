from test_base import TestBase
from src.document_processor import DocumentProcessor
import docx
import os

class TestDocumentProcessor(TestBase):
    def setUp(self):
        super().setUp()
        self.processor = DocumentProcessor(self.test_config)
        
        # Create test template
        doc = docx.Document()
        doc.add_paragraph("Hello [Name]!")
        doc.add_paragraph("PIN: [Pin]")
        doc.add_table(rows=1, cols=2)
        doc.tables[0].cell(0, 0).text = "Username:"
        doc.tables[0].cell(0, 1).text = "[Username]"
        
        self.template_path = os.path.join(self.test_templates_dir, "test.docx")
        doc.save(self.template_path)

    def test_process_document(self):
        output_path = os.path.join(self.test_output_dir, "output.docx")
        replacements = {
            "[Name]": "John Doe",
            "[Pin]": "123456",
            "[Username]": "johndoe"
        }
        
        success, message = self.processor.process_document(
            self.template_path,
            output_path,
            replacements
        )
        
        self.assertTrue(success)
        self.assertTrue(os.path.exists(output_path))
        
        # Verify replacements
        doc = docx.Document(output_path)
        self.assertEqual(doc.paragraphs[0].text, "Hello John Doe!")
        self.assertEqual(doc.paragraphs[1].text, "PIN: 123456")
        self.assertEqual(doc.tables[0].cell(0, 1).text, "johndoe") 