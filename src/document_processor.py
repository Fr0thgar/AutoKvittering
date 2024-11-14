import os
from datetime import datetime
import docx
import shutil
from utils.logger import Logger

class DocumentProcessor:
    def __init__(self, config):
        self.config = config
        self.logger = Logger()
        
    def process_document(self, template_path, output_path, replacements):
        try:
            # Create output directory if it doesn't exist
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Copy template to output location
            shutil.copyfile(template_path, output_path)
            
            # Process the document
            doc = docx.Document(output_path)
            replacements_made = self._process_replacements(doc, replacements)
            
            # Save the document
            doc.save(output_path)
            
            self.logger.info(f"Document processed successfully: {output_path}")
            return True, "Document created successfully"
            
        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")
            return False, str(e)
    
    def _process_replacements(self, doc, replacements):
        replacements_made = False
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            if self._replace_text(paragraph, replacements):
                replacements_made = True
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self._replace_text(paragraph, replacements):
                            replacements_made = True
        
        return replacements_made
    
    def _replace_text(self, paragraph, replacements):
        original_text = paragraph.text
        modified_text = original_text
        
        for placeholder, value in replacements.items():
            if placeholder in modified_text:
                modified_text = modified_text.replace(placeholder, value)
        
        if original_text != modified_text:
            paragraph.text = modified_text
            return True
        return False 