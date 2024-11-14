import json
from tkinter import *
import os
import glob
import shutil
import docx
import customtkinter
import random
import string
from datetime import *

customtkinter.set_default_color_theme("./theme.json")

# function to load config with file paths
def load_config():
    try:
        with open('config.json', 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        # default config if file doesnt exist
        default_config = {
            "template_directory": os.path.join(os.getcwd(), "templates"),
            "window_size": {"width": 500, "height": 400},
            "password_settings": {
                "length": 12,
                "use_uppercase": True,
                "use_lowercase": True,
                "use_digits": True,
                "use_special": True,
                "special_chars": "!@#$%^&*"
            }
        }
        # create config file with default values
        save_config(default_config)
        return default_config

def save_config(config):
    with open('config.json', 'w') as f:
        json.dump(config, indent=4)

def update_template_directory(new_path):
    config = load_config()
    config["template_directory"] = new_path
    save_config(config)

def browse_template_directory():
    from tkinter import filedialog
    config = load_config()
    new_dir = filedialog.askdirectory(initialdir=config["template_directory"])
    if new_dir:
        update_template_directory(new_dir)
        # Refresh template dropdown
        template_map = get_template_files()
        template_dropdown.configure(values=list(template_map.keys()))
        if template_map:
            template_var.set(list(template_map.keys())[0])

def get_template_files():
    config = load_config()
    template_dir = config["template_directory"]
    # Get full paths but only display filenames
    full_paths = glob.glob(os.path.join(template_dir, "*.docx"))
    # Create a dictionary mapping display names to full paths
    template_map = {os.path.basename(path): path for path in full_paths}
    return template_map

def generate_secure_password():
    config = load_config()
    password_settings = config["password_settings"]
    
    # Initialize character sets based on config
    chars = ""
    required_chars = []
    
    if password_settings["use_lowercase"]:
        chars += string.ascii_lowercase
        required_chars.append(random.choice(string.ascii_lowercase))
        
    if password_settings["use_uppercase"]:
        chars += string.ascii_uppercase
        required_chars.append(random.choice(string.ascii_uppercase))
        
    if password_settings["use_digits"]:
        chars += string.digits
        required_chars.append(random.choice(string.digits))
        
    if password_settings["use_special"]:
        special = password_settings["special_chars"]
        chars += special
        required_chars.append(random.choice(special))
    
    # Fill remaining length with random characters
    remaining_length = password_settings["length"] - len(required_chars)
    password = required_chars + [random.choice(chars) for _ in range(remaining_length)]
    
    # Shuffle the password
    random.shuffle(password)
    
    return ''.join(password)

def submit_name(event=None):
    name = name_entry.get()
    username = username_entry.get()
    selected_template_name = template_var.get()

    if name and username and selected_template_name: 
        # Generate PIN and password
        pin = str(random.randint(100000, 999999))
        password = generate_secure_password()  # This generates the password using config settings

        # Get current month and date 
        current_month = datetime.now().strftime("%B")
        current_date = datetime.now().strftime("%d")
        current_year = datetime.now().strftime("%Y")
        folder_path = os.path.join(os.getcwd(), current_year, current_month, current_date)
        os.makedirs(folder_path, exist_ok=True)

        # Create output path
        output_path = os.path.join(folder_path, f"{name}.docx")

        # Load the template document
        template_path = template_map[selected_template_name]
        shutil.copyfile(template_path, output_path)

        doc = docx.Document(output_path)

        # Define all placeholder types including password
        placeholders = ['[Name]', '[NAME]', '[name]','<Name>','<NAME>', '<name>']
        pin_placeholders = ['[Pin]', '[PIN]', '[pin]']
        username_placeholders = ['[Username]', '[USERNAME]', '[username]']
        password_placeholders = ['[Password]', '[PASSWORD]', '[password]']  # Password placeholders
        replacements_made = False

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            modified_text = original_text

            # Replace all types of placeholders
            for placeholder in placeholders:
                if placeholder in modified_text:
                    modified_text = modified_text.replace(placeholder, name)
                    replacements_made = True
            
            for placeholder in pin_placeholders:
                if placeholder in modified_text:
                    modified_text = modified_text.replace(placeholder, pin)
                    replacements_made = True
                    
            for placeholder in username_placeholders:
                if placeholder in modified_text:
                    modified_text = modified_text.replace(placeholder, username)
                    replacements_made = True
            
            # Replace password placeholders
            for placeholder in password_placeholders:
                if placeholder in modified_text:
                    modified_text = modified_text.replace(placeholder, password)
                    replacements_made = True
            
            if original_text != modified_text:
                paragraph.text = modified_text

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        modified_text = original_text

                        for placeholder in placeholders:
                            if placeholder in modified_text:
                                modified_text = modified_text.replace(placeholder, name)
                                replacements_made = True
                        
                        for placeholder in pin_placeholders:
                            if placeholder in modified_text:
                                modified_text = modified_text.replace(placeholder, pin)
                                replacements_made = True
                                
                        for placeholder in username_placeholders:
                            if placeholder in modified_text:
                                modified_text = modified_text.replace(placeholder, username)
                                replacements_made = True
                                
                        # Replace password placeholders in tables
                        for placeholder in password_placeholders:
                            if placeholder in modified_text:
                                modified_text = modified_text.replace(placeholder, password)
                                replacements_made = True
                        
                        if original_text != modified_text:
                            paragraph.text = modified_text

        # Save the modified document
        doc.save(output_path)
        
        if replacements_made:
            success_msg = f"Document created with replacements: {output_path}"
            success_msg += f"\nGenerated PIN: {pin}"
            print(success_msg)
        else:
            print(f"Warning: No placeholders found in template: {selected_template_name}")
            print("Expected placeholders:", placeholders)
            print("and:", pin_placeholders)
            print("and:", username_placeholders)
    else:
        print("Please enter a name, username, and select a template.")

    
# create the main app window
config = load_config()
root = customtkinter.CTk()
root.title("Insert into Word")
root.geometry(f"{config['window_size']['width']}x{config['window_size']['height']}")

# add template selection dropdown
template_label = customtkinter.CTkLabel(root, text="Select Template: ")
template_label.pack()

template_var = customtkinter.StringVar()
template_map = get_template_files()
template_dropdown = customtkinter.CTkOptionMenu(
    root,
    variable=template_var,
    values=list(template_map.keys())  # Only show filenames in dropdown
)
template_dropdown.pack()

settings_button = customtkinter.CTkButton(
    root,
    text="Change Template Directory",
    command=browse_template_directory
)
settings_button.pack(pady=10)

# If templates were found, set default value
if template_map:
    template_var.set(list(template_map.keys())[0])

# Create a label for the name input
name_label = customtkinter.CTkLabel(root, text="Enter Name: ")
name_label.pack()

# Create a text box for the name input
name_entry = customtkinter.CTkEntry(root)
name_entry.pack()

# Bind the Enter key to submit the name
name_entry.bind("<Return>", lambda event: submit_name())

# Add username label and entry after the name entry
username_label = customtkinter.CTkLabel(root, text="Enter Username: ")
username_label.pack()

username_entry = customtkinter.CTkEntry(root)
username_entry.pack()

# Bind Enter key to username entry as well
username_entry.bind("<Return>", lambda event: submit_name())

# Create a button to submit the name
submit_button = customtkinter.CTkButton(root, text="Submit", command=submit_name)
submit_button.pack()

root.mainloop()